import os
import tempfile
import docx
import fitz  # PyMuPDF
from fpdf import FPDF
from aiogram import Bot


# Форматирование вывода дат
def get_dates(list_dates: list):
    st = ''
    for idx, el in enumerate(list_dates):
        name = el.get('name', '')
        assignee = el.get('assignee', '')
        date = el.get('date', '')
        description = el.get('description', '')

        st += f"{idx + 1}. {name}\n"

        if assignee:
            st += f"Ответственный: {assignee}\n"
        else:
            st += "Ответственный не назначен\n"

        st += f"Назначенная дата: {date}\nОписание: {description}\n\n"

    return st.strip()


# Функция парсинга DOCX из файла
def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])


# Функция парсинга PDF из файла
def extract_text_from_pdf(file_path: str) -> str:
    text = []
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text.append(page.get_text())
    return '\n'.join(text)


# Конвертирует файл в текст (PDF, DOCX, TXT)
async def file_to_text(bot: Bot, file_id: str) -> str:
    file = await bot.get_file(file_id)
    ext = os.path.splitext(file.file_path)[1].lower()

    if ext not in ('.pdf', '.docx', '.txt'):
        raise ValueError("Неподдерживаемый формат файла")

    # Временный путь
    tmp_dir = tempfile.gettempdir()
    tmp_path = os.path.join(tmp_dir, os.path.basename(file.file_path))

    # Скачиваем файл
    await bot.download_file(file.file_path, tmp_path)

    try:
        if ext == '.pdf':
            return extract_text_from_pdf(tmp_path)
        elif ext == '.docx':
            return extract_text_from_docx(tmp_path)
        else:
            with open(tmp_path, 'r', encoding='utf-8') as f:
                return f.read()
    finally:
        os.remove(tmp_path)


# Сохраняет текст в .docx и возвращает путь к файлу
def text_to_docx(text: str) -> str:
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)

    file_path = os.path.join(tempfile.gettempdir(), "meet_flow.docx")
    doc.save(file_path)
    return file_path


# Сохраняет текст в PDF с поддержкой кириллицы, используя встроенный шрифт FreeSerif.
def text_to_pdf(text: str) -> str:
    file_path = os.path.join(tempfile.gettempdir(), "meet_flow.pdf")
    pdf = FPDF()
    pdf.add_page()

    font_path = "fonts/FreeSerif.ttf"
    if not os.path.exists(font_path):
        raise FileNotFoundError("Файл шрифта fonts/FreeSerif.ttf не найден.")

    # Регистрируем шрифт
    pdf.add_font("FreeSerif", "", font_path, uni=True)
    pdf.set_font("FreeSerif", "", 12)

    margin_left = 15  # мм
    width_mm = 180  # ширина блока (210 мм минус поля)
    line_height = 8

    pdf.set_left_margin(margin_left)
    pdf.set_right_margin(15)
    pdf.set_auto_page_break(auto=True, margin=20)

    for line in text.split('\n'):
        pdf.set_x(margin_left)  # ВАЖНО: сброс X перед каждой строкой
        if line.strip():
            pdf.multi_cell(width_mm, line_height, line)
        else:
            pdf.ln(line_height)

    pdf.output(file_path)
    return file_path
